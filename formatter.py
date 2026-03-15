"""
Hedda & Pico Skript-Formatter — hardcoded Formatierung für Hörspiel-Manuskripte.

Tested against all 7 Folgen (01–06, 12) with varying input formats:
- Some have empty lines between every paragraph, others have none
- Some use [square brackets], others use (round) — all get normalized
- Some have SFX/ATM prefix, others have bare [description] for sounds
- Folge 04 starts with empty paragraph + uses "Text A" style
- PICO-STOP is a special direction type
- LEIT-OBJEKT has many variants (EINFÜHRUNG, HÖHEPUNKT, WENDEPUNKT, etc.)

Two-pass approach:
1. Classify every paragraph
2. Format based on classification + speaker context
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import io
from datetime import datetime


# ─── Hardcoded rules ─────────────────────────────────────────────────────────

FONT = "Calibri"
BODY_SIZE = 12
LINE_SPACING = 1.5
NAME_LINE_SPACING = 1.0       # single spacing on name line (0.5 clips text in Word Desktop)
STAGE_DIR_SIZE = 9             # Regieanweisungen, SFX, ATM, PICO-STOP
TITLE_SIZE = 48                # "HEDDA & PICO"
EPISODE_SIZE = 26              # "Folge X: ..."
SCENE_SIZE = 13                # "SZENE 1: ..."
TIME_SIZE = 9                  # kumulierte Zeit
PAGE_NUM_SIZE = 9

# Character colors (name_color only — dialogue text stays black unless specified)
CHARACTER_COLORS = {
    "ERZÄHLER":     {"name_color": "#000000", "text_italic": True, "text_indent_cm": 1.27},
    "WENDT":        {"name_color": "#00B050"},
    "HEDDA":        {"name_color": "#FF8C00"},
    "FRAU FISCHER": {"name_color": "#C88A00"},
    "HERR NOVAK":   {"name_color": "#008B8B"},
    "HERR HASSAN":  {"name_color": "#8B4513"},
    "OMA STEIN":    {"name_color": "#FF69B4"},
    "DR. KHOURY":   {"name_color": "#FF0000"},
}

CASE_CHARACTER_COLOR = "#800080"  # lila
BLUE_SHADES = ["#0000FF", "#00008B", "#4169E1", "#1E90FF", "#00BFFF", "#87CEEB"]

# Prefixes that are NEVER character names
NOT_CHARACTER_PREFIXES = (
    "SZENE", "SFX", "ATM", "LEIT", "OUTRO", "TITELSONG", "PICO-STOP",
    "HEDDA &",  # title line
)


# ─── Helpers ──────────────────────────────────────────────────────────────────

def _hex_to_rgb(hex_color):
    h = hex_color.lstrip("#")
    try:
        return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))
    except (ValueError, IndexError):
        return RGBColor(0, 0, 0)


def _set_font(run, font_name=FONT):
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = run._element.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rFonts.set(qn(f"w:{attr}"), font_name)


def _fmt_run(run, size=BODY_SIZE, bold=False, italic=False, color="#000000", underline=False):
    _set_font(run)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    run.font.color.rgb = _hex_to_rgb(color)
    run.font.highlight_color = None


def _fmt_para(para, align="left", spacing=LINE_SPACING, indent_cm=None,
              space_before=None, space_after=None):
    pf = para.paragraph_format
    pf.line_spacing = spacing
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    align_map = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                 "right": WD_ALIGN_PARAGRAPH.RIGHT, "justify": WD_ALIGN_PARAGRAPH.JUSTIFY}
    if align in align_map:
        pf.alignment = align_map[align]
    if indent_cm is not None:
        pf.left_indent = Cm(indent_cm)
    if space_before is not None:
        pf.space_before = Pt(space_before)
    if space_after is not None:
        pf.space_after = Pt(space_after)


def _add_page_numbers(doc):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    run = p.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    run._element.append(fld)

    run2 = p.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run2._element.append(instr)

    run3 = p.add_run()
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run3._element.append(fld2)

    for r in (run, run2, run3):
        r.font.size = Pt(PAGE_NUM_SIZE)
        _set_font(r)
        r.font.color.rgb = RGBColor(0, 0, 0)


def _brackets_to_round(text):
    return text.replace("[", "(").replace("]", ")")


# ─── Classification ───────────────────────────────────────────────────────────

T_EMPTY = "empty"
T_TITLE = "title"
T_EPISODE = "episode"
T_SUBTITLE = "subtitle"   # "30-Minuten Kinder-Hörspiel-Manuskript"
T_SCENE = "scene"
T_CHARACTER = "character"
T_DIALOGUE = "dialogue"
T_STAGE_DIR = "stage_dir"
T_SFX_ATM = "sfx_atm"     # SFX, ATM, TITELSONG, OUTRO, PICO-STOP
T_TIME = "time"
T_LEIT = "leit"
T_BODY = "body"


def _known_names():
    return {n.upper() for n in CHARACTER_COLORS}


def _is_character(text, known):
    """Check if text is a standalone character name line."""
    s = text.strip()
    if not s or len(s.split()) > 5:
        return False

    upper = s.upper()

    # Reject known non-character prefixes
    for prefix in NOT_CHARACTER_PREFIXES:
        if upper.startswith(prefix):
            return False

    # Exact match against known characters
    if upper in known:
        return True

    # Heuristic: all uppercase, short, no special punctuation
    # Handles unknown characters like "KINDER", "JONAS", "LINA" etc.
    clean = s.replace(".", "").replace(" ", "").replace("-", "")
    if clean and clean == clean.upper() and clean.isalpha():
        return True

    return False


def _classify(texts):
    """
    Classify every paragraph. Handles all known Folge formats:
    - With/without empty lines between paragraphs
    - Square brackets and round brackets
    - With/without SFX/ATM prefix
    - Leading empty paragraphs (Folge 04)
    - Subtitle line (Folge 04)
    """
    known = _known_names()
    result = []
    title_seen = False
    episode_seen = False

    for text in texts:
        s = text.strip()

        # Empty
        if not s:
            result.append((T_EMPTY, None))
            continue

        # Title — first line containing "Hedda" and "Pico"
        if not title_seen and re.search(r"Hedda.*Pico|HEDDA.*PICO", s, re.IGNORECASE):
            title_seen = True
            result.append((T_TITLE, None))
            continue

        # Episode — "Folge N: ..."
        if not episode_seen and re.match(r".*Folge\s+\d+", s, re.IGNORECASE):
            episode_seen = True
            result.append((T_EPISODE, None))
            continue

        # Subtitle — "30-Minuten Kinder-Hörspiel-Manuskript" (only Folge 04)
        if title_seen and not episode_seen and re.match(r"\d+-Minuten", s, re.IGNORECASE):
            result.append((T_SUBTITLE, None))
            continue

        # Time marker — [Kumulierte Zeit: ...] or (Kumulierte Zeit: ...)
        if re.match(r"^[\[\(]?\s*Kumulierte\s+Zeit", s, re.IGNORECASE):
            result.append((T_TIME, None))
            continue

        # Scene heading — SZENE N: ...
        if re.match(r"^SZENE\s+\d+|^Szene\s+\d+", s):
            result.append((T_SCENE, None))
            continue

        # SFX/ATM/TITELSONG/OUTRO/PICO-STOP — with prefix in brackets
        if re.match(r"^[\[\(]\s*(SFX|ATM|TITELSONG|OUTRO|PICO-STOP)", s, re.IGNORECASE):
            result.append((T_SFX_ATM, None))
            continue

        # LEIT-OBJEKT — all variants (EINFÜHRUNG, HÖHEPUNKT, WENDEPUNKT, etc.)
        if re.search(r"LEIT-OBJEKT", s, re.IGNORECASE):
            result.append((T_LEIT, None))
            continue

        # Full-line brackets/parens — stage directions or sound descriptions
        if (s.startswith("[") and s.endswith("]")) or \
           (s.startswith("(") and s.endswith(")")):
            result.append((T_STAGE_DIR, None))
            continue

        # Character name (standalone line, uppercase, short)
        if _is_character(s, known):
            result.append((T_CHARACTER, s.upper()))
            continue

        # Default: body text
        result.append((T_BODY, None))

    # ── Second pass: assign speaker context ──
    final = []
    speaker = None
    for ptype, data in result:
        if ptype == T_CHARACTER:
            speaker = data
            final.append((T_CHARACTER, data))
        elif ptype == T_EMPTY:
            # Empty lines keep current speaker context
            final.append((T_EMPTY, speaker))
        elif ptype in (T_SCENE, T_TITLE, T_EPISODE, T_SUBTITLE):
            speaker = None
            final.append((ptype, None))
        elif ptype == T_BODY:
            # Body text after a character = dialogue
            final.append((T_DIALOGUE, speaker) if speaker else (T_BODY, None))
        elif ptype == T_STAGE_DIR:
            final.append((T_STAGE_DIR, speaker))
        else:
            final.append((ptype, speaker))

    return final


# ─── Character color lookup ───────────────────────────────────────────────────

def _char_config(name, unknown_tracker, case_characters):
    """Get color config for a character. Case characters get purple."""
    upper = name.upper().strip()

    for known, cfg in CHARACTER_COLORS.items():
        if known == upper:
            return cfg

    if upper not in unknown_tracker:
        if upper in case_characters:
            unknown_tracker[upper] = {"name_color": CASE_CHARACTER_COLOR}
        else:
            blue_count = sum(1 for v in unknown_tracker.values()
                            if v.get("name_color") != CASE_CHARACTER_COLOR)
            idx = blue_count % len(BLUE_SHADES)
            unknown_tracker[upper] = {"name_color": BLUE_SHADES[idx]}

    return unknown_tracker.get(upper, {"name_color": "#0000FF"})


# ─── Main formatting ─────────────────────────────────────────────────────────

def format_document(source_path, case_characters=None):
    """
    Format a Hedda & Pico script. Returns .docx bytes.
    case_characters: set of character names (uppercase) that get purple color.
    """
    if case_characters is None:
        case_characters = set()
    else:
        case_characters = {c.upper().strip() for c in case_characters}

    source = Document(source_path)
    texts = [p.text for p in source.paragraphs]
    classifications = _classify(texts)

    # Pre-pass: find stage directions to merge onto character name line
    # A stage direction immediately after a character name (possibly with
    # empty lines between, as in Folge 06/12) gets merged onto the name line
    merge_set = set()
    for idx, (pt, _) in enumerate(classifications):
        if pt == T_CHARACTER:
            nxt = idx + 1
            while nxt < len(classifications) and classifications[nxt][0] == T_EMPTY:
                nxt += 1
            if nxt < len(classifications) and classifications[nxt][0] == T_STAGE_DIR:
                merge_set.add(nxt)

    # New document
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(BODY_SIZE)
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.line_spacing = LINE_SPACING
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    _add_page_numbers(doc)
    unknown = {}

    for i, (ptype, speaker) in enumerate(classifications):
        text = _brackets_to_round(texts[i].strip())

        # Skip merged stage directions (already on character name line)
        if i in merge_set:
            continue

        # ── EMPTY ──
        # All empty lines between content paragraphs are suppressed.
        # Spacing is fully controlled by space_before/space_after.
        # Only exception: one blank line after title+episode block (before first scene).
        if ptype == T_EMPTY:
            # Keep one blank line between episode/subtitle and first scene
            prev = i - 1
            while prev >= 0 and classifications[prev][0] == T_EMPTY:
                prev -= 1
            nxt = i + 1
            while nxt < len(classifications) and classifications[nxt][0] == T_EMPTY:
                nxt += 1

            # Blank between title block (episode/subtitle/body) and scene
            if prev >= 0 and classifications[prev][0] in (T_EPISODE, T_SUBTITLE, T_BODY) \
               and nxt < len(classifications) and classifications[nxt][0] == T_SCENE:
                # Only emit ONE blank, skip duplicates
                if i == prev + 1:  # first blank in a run
                    p = doc.add_paragraph()
                    _fmt_para(p)
                continue

            # All other blanks: suppress
            continue

        # ── TITLE ──
        if ptype == T_TITLE:
            p = doc.add_paragraph()
            _fmt_para(p)
            run = p.add_run(text.upper())
            _fmt_run(run, size=TITLE_SIZE, bold=True)
            continue

        # ── EPISODE ──
        if ptype == T_EPISODE:
            p = doc.add_paragraph()
            _fmt_para(p, space_before=6)
            run = p.add_run(text)
            _fmt_run(run, size=EPISODE_SIZE, bold=True)
            continue

        # ── SUBTITLE (e.g., "30-Minuten Kinder-Hörspiel-Manuskript") ──
        if ptype == T_SUBTITLE:
            p = doc.add_paragraph()
            _fmt_para(p)
            run = p.add_run(text)
            _fmt_run(run)
            continue

        # ── SCENE ──
        if ptype == T_SCENE:
            # Add one blank line before scene, but check if one was already
            # emitted by the EMPTY handler (title block → scene transition)
            last_para_text = doc.paragraphs[-1].text.strip() if doc.paragraphs else ""
            if i > 0 and last_para_text:
                bp = doc.add_paragraph()
                _fmt_para(bp)
            p = doc.add_paragraph()
            _fmt_para(p)
            run = p.add_run(text.upper())
            _fmt_run(run, size=SCENE_SIZE, bold=True, underline=True)
            continue

        # ── TIME MARKER ──
        if ptype == T_TIME:
            p = doc.add_paragraph()
            _fmt_para(p, align="right")
            run = p.add_run(text)
            _fmt_run(run, size=TIME_SIZE, italic=True)
            continue

        # ── CHARACTER NAME ──
        if ptype == T_CHARACTER:
            cfg = _char_config(speaker, unknown, case_characters)
            name_color = cfg.get("name_color", "#000000")

            p = doc.add_paragraph()
            _fmt_para(p, spacing=NAME_LINE_SPACING, space_before=6, space_after=0)
            run = p.add_run(text.upper())
            _fmt_run(run, bold=True, color=name_color)

            # Merge following stage direction onto same line
            nxt = i + 1
            while nxt < len(classifications) and classifications[nxt][0] == T_EMPTY:
                nxt += 1
            if nxt in merge_set:
                dir_text = _brackets_to_round(texts[nxt].strip())
                sp = p.add_run(" ")
                _fmt_run(sp, color=name_color)
                dr = p.add_run(dir_text)
                _fmt_run(dr, size=STAGE_DIR_SIZE, italic=True)
            continue

        # ── DIALOGUE ──
        if ptype == T_DIALOGUE and speaker:
            cfg = _char_config(speaker, unknown, case_characters)
            text_color = cfg.get("text_color", "#000000")
            text_italic = cfg.get("text_italic", False)
            indent = cfg.get("text_indent_cm", None)

            p = doc.add_paragraph()
            _fmt_para(p, indent_cm=indent, space_after=0)
            run = p.add_run(text)
            _fmt_run(run, italic=text_italic, color=text_color)
            continue

        # ── STAGE DIRECTION (standalone, on its own line) ──
        if ptype == T_STAGE_DIR:
            indent = None
            if speaker:
                cfg = _char_config(speaker, unknown, case_characters)
                indent = cfg.get("text_indent_cm", None)
            p = doc.add_paragraph()
            _fmt_para(p, indent_cm=indent, space_after=0)
            run = p.add_run(text)
            _fmt_run(run, size=STAGE_DIR_SIZE, italic=True)
            continue

        # ── SFX/ATM/TITELSONG/OUTRO/PICO-STOP ──
        if ptype == T_SFX_ATM:
            p = doc.add_paragraph()
            _fmt_para(p, space_after=0)
            run = p.add_run(text)
            _fmt_run(run, size=STAGE_DIR_SIZE, italic=True)
            continue

        # ── LEIT-OBJEKT ──
        if ptype == T_LEIT:
            p = doc.add_paragraph()
            _fmt_para(p, space_after=0)
            # Find "LEIT-OBJEKT" in text and make it bold
            m = re.search(r"LEIT-OBJEKT", text, re.IGNORECASE)
            if m:
                before = text[:m.start()]
                leit = text[m.start():m.end()]
                after = text[m.end():]
                if before:
                    r = p.add_run(before)
                    _fmt_run(r, size=STAGE_DIR_SIZE, italic=True)
                r = p.add_run(leit.upper())
                _fmt_run(r, size=STAGE_DIR_SIZE, bold=True, italic=True)
                if after:
                    r = p.add_run(after)
                    _fmt_run(r, size=STAGE_DIR_SIZE, italic=True)
            else:
                r = p.add_run(text)
                _fmt_run(r, size=STAGE_DIR_SIZE, italic=True)
            continue

        # ── DEFAULT BODY ──
        p = doc.add_paragraph()
        _fmt_para(p, space_after=0)
        run = p.add_run(text)
        _fmt_run(run)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def get_output_filename(original_name):
    """Generate output filename with date suffix (DDMM)."""
    base = original_name.replace(".docx", "")
    suffix = datetime.now().strftime("%d%m")
    return f"{base}_{suffix}.docx"
